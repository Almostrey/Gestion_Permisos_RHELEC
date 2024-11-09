# System
from sys import argv
from os import startfile, getcwd, mkdir, remove
from os.path import isfile
from shutil import rmtree
# PySide6
from PySide6 import QtSvg
from PySide6 import QtCore
from PySide6 import QtWidgets
from PySide6 import QtGui
from PySide6 import QtMultimedia
# CSV, Hashlib, datetime, openpyxl, docx, pandas
from pandas import read_excel, read_table, DataFrame, ExcelWriter
from hashlib import sha256
from csv import reader
from threading import Timer
from datetime import datetime, timedelta
from openpyxl import load_workbook
from docx import Document
import msoffcrypto
from io import BytesIO
from pyperclip import copy
# QT_Windows
import QT_Windows
import QT_Windows.WindowSelectEmail_ui
import QT_Windows.WindowLogin_ui
import QT_Windows.WindowSolicitudPermisos_ui
import QT_Windows.VideoLogoWindow_ui
# Email Libreries
from imaplib import IMAP4_SSL
from email import message_from_bytes
global usuario
usuario = ""
class WindowSendEmail(QtWidgets.QMainWindow, QT_Windows.WindowSelectEmail_ui.Ui_SeleccionCorreo):
    def __init__(self, allData:dict):
        super(WindowSendEmail, self).__init__()
        self.username = "deguevarab@hotmail.com"
        self.password = "auplxlzlvuukuhqn"
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.FramelessWindowHint)
        self.setAttribute(QtCore.Qt.WA_TranslucentBackground)
        self.allData = allData
        self.pbClose.clicked.connect(self.closeWindow)
        self.pbMaximize.clicked.connect(self.maximizeWindow)
        self.pbMinimize.clicked.connect(self.showMinimized)
        # Clave auplxlzlvuukuhqn
        '''self.loadPreviousEmails()
        self.loadBodyEmail()'''
    def loadPreviousEmails(self):
        #imap_server = IMAP4_SSL('smtp.gmail.com')
        imap_server = IMAP4_SSL('outlook.office365.com', 993)
        imap_server.login(self.username, self.password)

        status, messages = imap_server.select('INBOX')
        result, data = imap_server.search(None, f'SUBJECT "{self.allData["NameRBS"]}"')
        print("IDS: "+str(data))
        result, data = imap_server.fetch("615", "(RFC822)")
        a = data[0][1].split()
        raw_email = data[0][1]
        email_message = message_from_bytes(raw_email)

        for i in range(int(messages[0]), int(messages[0])-5, -1):
            res, msg = imap_server.fetch(str(i), "(RFC822)")
            print(msg)
        pass
    def loadBodyEmail(self):
        pass
    def maximizeWindow(self):
        if not self.isMaximized():self.showMaximized()
        else: self.showNormal()
    def closeWindow(self):
        self.close()
    def mouseMoveEvent(self, event):
        if self.oldPosition.y()-self.y() <= 45 and not self.isMaximized():
            p = event.globalPosition()
            delta = QtCore.QPoint(p.toPoint()-self.oldPosition)
            self.move(self.x()+delta.x(), self.y()+delta.y())
            p = event.globalPosition()
            self.oldPosition = p.toPoint()
    def mousePressEvent(self, event):
        p = event.globalPosition()
        self.oldPosition = p.toPoint()

class WindowLogin(QtWidgets.QMainWindow, QT_Windows.WindowLogin_ui.Ui_LoginWindow):
    def __init__(self):
        super(WindowLogin, self).__init__()
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.FramelessWindowHint)
        self.setAttribute(QtCore.Qt.WA_TranslucentBackground)
        self.pbClose.clicked.connect(self.closeWindow)
        self.pbMinimize.clicked.connect(self.showMinimized)
        self.pbAcceder.clicked.connect(self.login)
    def closeWindow(self):
        deleteAllData()
        self.close()
    def mouseMoveEvent(self, event):
        if self.oldPosition.y()-self.y() <= 45 and not self.isMaximized():
            p = event.globalPosition()
            delta = QtCore.QPoint(p.toPoint()-self.oldPosition)
            self.move(self.x()+delta.x(), self.y()+delta.y())
            p = event.globalPosition()
            self.oldPosition = p.toPoint()
    def mousePressEvent(self, event):
        p = event.globalPosition()
        self.oldPosition = p.toPoint()
    def keyPressEvent(self, event):
        if event.key() == QtCore.Qt.Key_Return or event.key() == QtCore.Qt.Key_Enter:
            if self.fieldsFilled():
                self.login()
            else: pass
        else: pass
    def normalFields(self):
        self.frameUsername.setStyleSheet("""border-radius:15px;color: rgb(0, 0, 0);background-color: rgb(190, 190, 190);border:0px;""")
        self.framePassword.setStyleSheet("""border-radius:15px;color: rgb(0, 0, 0);background-color: rgb(190, 190, 190);border:0px;""")
        self.lblIconUsername.setStyleSheet("""image: url(:/featherIcons/ResourcesFolder/featherIcons/user.svg);background-color: rgb(190, 190, 190);border-radius: 10px;border:0px;""")
        self.lblIconPassword.setStyleSheet("""image: url(:/featherIcons/ResourcesFolder/featherIcons/lock.svg);background-color: rgb(190, 190, 190);border-radius: 10px;border:0px;""")
        self.txtUsername.setStyleSheet("""border-radius:15px;color: rgb(0, 0, 0);background-color: rgb(190, 190, 190);border:0px;""")
        self.txtPassword.setStyleSheet("""border-radius:15px;color: rgb(0, 0, 0);background-color: rgb(190, 190, 190);border:0px;""")
        #self.txtUsername.setText(self.txtUsername.text())
        #self.txtPassword.setText(self.txtPassword.text())
        self.pbAcceder.setStyleSheet("""QPushButton{border-radius:15px;color: rgb(0, 0, 0);background-color: rgb(154, 192, 213);border:0px;}
                                     QPushButton:hover{background-color: rgb(124, 162, 183);}
                                     QPushButton:pressed{background-color: rgb(154, 192, 213);}""")
        self.pbAcceder.setText("Acceder")
    def fieldsFilled(self):
        if (self.txtUsername.text() != "" and self.txtPassword.text() != ""): return True
        else: 
            if self.txtUsername.text() != "":
                self.framePassword.setStyleSheet("""border-radius:15px;color: rgb(0, 0, 0);background-color: rgb(226, 102, 0);border:0px;""")
                self.txtPassword.setStyleSheet("""border-radius:15px;color: rgb(0, 0, 0);background-color: rgb(226, 102, 0);border:0px;""")
                self.lblIconPassword.setStyleSheet("""image: url(:/featherIcons/ResourcesFolder/featherIcons/lock.svg);background-color: rgb(226, 102, 0);border-radius: 10px;border:0px;""")
                self.pbAcceder.setText("Ingrese clave")
            else: pass
            if self.txtPassword.text() != "":
                self.frameUsername.setStyleSheet("""border-radius:15px;color: rgb(0, 0, 0);background-color: rgb(226, 102, 0);border:0px;""")
                self.txtUsername.setStyleSheet("""border-radius:15px;color: rgb(0, 0, 0);background-color: rgb(226, 102, 0);border:0px;""")
                self.lblIconUsername.setStyleSheet("""image: url(:/featherIcons/ResourcesFolder/featherIcons/user.svg);background-color: rgb(226, 102, 0);border-radius: 10px;border:0px;""")
                self.pbAcceder.setText("Ingrese usuario")
            else: pass
            Timer(2, self.normalFields).start()
            return False
    def login(self):
        if self.fieldsFilled():
            #### GUI
            # Username GUI
            self.hash_object = sha256()
            self.hash_object.update(self.txtUsername.text().encode())
            self.username = self.hash_object.hexdigest()
            # Pass GUI
            self.hash_object = sha256()
            self.hash_object.update(self.txtPassword.text().encode())
            self.password = self.hash_object.hexdigest()
            #### Comparing with CSV
            with open("data/Login.csv") as self.csvLogin:
                self.csvInfoLogin = reader(self.csvLogin)
                for self.row in self.csvInfoLogin:
                    if str(self.row[0]).split(";")[4] == self.username:
                        if str(self.row[0]).split(";")[5] == self.password:
                            global usuario
                            usuario = str(self.row[0]).split(";")[1] +" "+ str(self.row[0]).split(";")[2]
                            self.close()
                            self.w = WindowSolicitudPermisos()
                            self.w.show()
                        else: pass
                    else: pass
            self.credentialsIncorrect()
        else: pass
    def credentialsIncorrect(self):
        self.usernameFound = False
        with open("data/Login.csv") as self.csvLogin:
            self.csvInfoLogin = reader(self.csvLogin)
            for self.row in self.csvInfoLogin:
                if str(self.row[0]).split(";")[4] == self.username: 
                    self.usernameFound = True
                    break
                else: pass
        if self.usernameFound:
            self.framePassword.setStyleSheet("""border-radius:15px;color: rgb(0, 0, 0);background-color: rgb(226, 102, 0);border:0px;""")
            self.txtPassword.setStyleSheet("""border-radius:15px;color: rgb(0, 0, 0);background-color: rgb(226, 102, 0);border:0px;""")
            self.lblIconPassword.setStyleSheet("""image: url(:/featherIcons/ResourcesFolder/featherIcons/lock.svg);background-color: rgb(226, 102, 0);border-radius: 10px;border:0px;""")
            self.txtPassword.setText("")
            self.pbAcceder.setText("Clave Incorrecta")
        else:
            self.frameUsername.setStyleSheet("""border-radius:15px;color: rgb(0, 0, 0);background-color: rgb(226, 102, 0);border:0px;""")
            self.txtUsername.setStyleSheet("""border-radius:15px;color: rgb(0, 0, 0);background-color: rgb(226, 102, 0);border:0px;""")
            self.lblIconUsername.setStyleSheet("""image: url(:/featherIcons/ResourcesFolder/featherIcons/user.svg);background-color: rgb(226, 102, 0);border-radius: 10px;border:0px;""")
            self.pbAcceder.setText("Usuario Incorrecto")
        Timer(2, self.normalFields).start()

class ReadOnlyDelegate(QtWidgets.QStyledItemDelegate):
    def createEditor(self, parent, option, index): return

class AlignDelegate(QtWidgets.QStyledItemDelegate):
    def initStyleOption(self, option, index):
        super(AlignDelegate, self).initStyleOption(option, index)
        option.displayAlignment = QtCore.Qt.AlignCenter

class WindowSolicitudPermisos(QtWidgets.QMainWindow, QT_Windows.WindowSolicitudPermisos_ui.Ui_SolicitudPermisos):
    def __init__(self):
        # Basic Setup and signals
        super(WindowSolicitudPermisos, self).__init__()
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.FramelessWindowHint)
        self.setAttribute(QtCore.Qt.WA_TranslucentBackground)
        self.pbClose.clicked.connect(self.closeWindow)
        self.pbMinimize.clicked.connect(self.showMinimized)
        self.pbMaximize.clicked.connect(self.maximizeWindow)
        self.cbRBS.currentIndexChanged.connect(self.loadDataFromRBS)
        self.cbCuadrillas.currentIndexChanged.connect(self.loadSupervisorNTableCuadrilla)
        self.cbOyMs.currentIndexChanged.connect(self.updateOyMTable)
        self.pbAgregarPermiso.clicked.connect(self.loadPermisosTable)
        self.pbEliminarPermiso.clicked.connect(self.eliminarPermisoTable)
        self.pbAgregarTecnico.clicked.connect(self.addPersonalTecnicoTable)
        self.pbEliminarTecnico.clicked.connect(self.eliminarPersonalTecnicoTable)
        self.pbEnviarFormatos.clicked.connect(self.enviarFormato)
        self.pbCopyCorreos.clicked.connect(self.copyCorreos)
        self.pbCopyOyMCorreo.clicked.connect(self.copyOyMCorreo)
        self.pbCopyOyMTel.clicked.connect(self.copyOyMTel)
        self.pbCopySupervisorCorreo.clicked.connect(self.copySupervisorCorreo)
        self.pbCopySupervisorTel.clicked.connect(self.copySupervisorTel)
        # init Functions
        self.tablePermisosNecesarios.setGeometry(self.tablePermisosNecesarios.x(), self.tablePermisosNecesarios.x(), 639, 191)
        self.tableTecnicosCuadrilla.setGeometry(self.tableTecnicosCuadrilla.x(), self.tableTecnicosCuadrilla.x(), 639, 192)
        self.adjustTablesSize()
        self.loadRBS()
        self.loadcbPermisos()
        self.fixHeaderTables()
    def copySupervisorCorreo(self):
        with open("data/Supervisores.csv") as self.csvSupervisores:
            self.InfoSupervisores = reader(self.csvSupervisores)
            for self.row in self.InfoSupervisores:
                if str(self.row[0]).split(";")[1] == self.cbSupervisores.currentText():
                    copy(str(self.row[0]).split(";")[3])
    def copySupervisorTel(self):
        with open("data/Supervisores.csv") as self.csvSupervisores:
            self.InfoSupervisores = reader(self.csvSupervisores)
            for self.row in self.InfoSupervisores:
                if str(self.row[0]).split(";")[1] == self.cbSupervisores.currentText():
                    copy(str(self.row[0]).split(";")[2])
    def copyCorreos(self):
        if self.lineRegion.text() == "R1":
            if self.getProvinciaCantonRBS()[1].upper() == "QUITO":
                if self.lineZona.text() == "Z1":
                    # R1Z1 UIO
                    # tecomreddeaccesor1z1@claro.com.ec; nohemi.guadalupe@rhelec.ec; kevin.moreno@rhelec.ec; Azambrano@claro.com.ec; ernesto.quimbita@rhelec.ec; daniel.guisha@rhelec.ec; rodriguezlem@globalhitss.com; mariajose.orozco@rhelec.ec;maritza.nunez@rhelec.ec
                    pass
                else:
                    # R1Z2 UIO
                    # tecoymrar1z2@claro.com.ec;  lgarcia@claro.com.ec; nohemi.guadalupe@rhelec.ec; kevin.moreno@rhelec.ec; Azambrano@claro.com.ec; ernesto.quimbita@rhelec.ec; daniel.guisha@rhelec.ec; rodriguezlem@globalhitss.com; mariajose.orozco@rhelec.ec; maritza.nunez@rhelec.ec
                    pass
            else:
                if self.lineZona.text() == "Z1":
                    # R1Z1 Prov
                    # tecomreddeaccesor1z1@claro.com.ec ; Azambrano@claro.com.ec; nohemi.guadalupe@rhelec.ec; kevin.moreno@rhelec.ec; ernesto.quimbita@rhelec.ec; daniel.guisha@rhelec.ec; rodriguezlem@globalhitss.com; mariajose.orozco@rhelec.ec; maritza.nunez@rhelec.ec
                    pass
                else:
                    # R1Z2 Prov
                    # tecoymrar1z2@claro.com.ec;  lgarcia@claro.com.ec; nohemi.guadalupe@rhelec.ec; kevin.moreno@rhelec.ec; Azambrano@claro.com.ec; ernesto.quimbita@rhelec.ec; daniel.guisha@rhelec.ec; rodriguezlem@globalhitss.com; mariajose.orozco@rhelec.ec; maritza.nunez@rhelec.ec
                    pass
        elif self.lineRegion.text() == "R2":
            if self.getProvinciaCantonRBS()[1].upper() == "GUAYAQUIL" or self.getProvinciaCantonRBS()[1].upper() == "SANTA ELENA" or self.getProvinciaCantonRBS()[1].upper() == "LA TRONCAL":
                if self.lineZona.text() == "Z1":
                    # R2Z1 GYE
                    # tecoymrar2z1@claro.com.ec; Azambrano@claro.com.ec; nohemi.guadalupe@rhelec.ec; kevin.moreno@rhelec.ec; alexander.santander@rhelec.ec;  stalin.tutiven@rhelec.ec; christian.navas@rhelec.ec; rodriguezlem@globalhitss.com; mariajose.orozco@rhelec.ec;
                    pass
                else:
                    # R2Z2 GYE
                    # tecoymrar2z2@claro.com.ec; Azambrano@claro.com.ec; nohemi.guadalupe@rhelec.ec; kevin.moreno@rhelec.ec; alexander.santander@rhelec.ec; stalin.tutiven@rhelec.ec; christian.navas@rhelec.ec; rodriguezlem@globalhitss.com; mariajose.orozco@rhelec.ec;
                    pass
            else:
                if self.lineZona.text() == "Z1":
                    # R2Z1 Prov
                    # tecoymrar2z1@claro.com.ec; Azambrano@claro.com.ec; nohemi.guadalupe@rhelec.ec; kevin.moreno@rhelec.ec; alexander.santander@rhelec.ec; stalin.tutiven@rhelec.ec; rodriguezlem@globalhitss.com;  christian.navas@rhelec.ec; mariajose.orozco@rhelec.ec;
                    pass
                else:
                    # R2Z2 Prov
                    # tecoymrar2z2@claro.com.ec; azambrano@claro.com.ec; nohemi.guadalupe@rhelec.ec; kevin.moreno@rhelec.ec; alexander.santander@rhelec.ec;  stalin.tutiven@rhelec.ec; christian.navas@rhelec.ec; rodriguezlem@globalhitss.com; mariajose.orozco@rhelec.ec;
                    pass
        else: pass
        pass
    def copyOyMCorreo(self):
        with open("data/OyMs.csv") as self.csvOYMs:
            self.InfoOYMs = reader(self.csvOYMs)
            for self.row in self.InfoOYMs:
                if str(self.row[0]).split(";")[1] == self.cbOyMs.currentText():
                    copy(str(self.row[0]).split(";")[3])
    def copyOyMTel(self):
        with open("data/OyMs.csv") as self.csvOYMs:
            self.InfoOYMs = reader(self.csvOYMs)
            for self.row in self.InfoOYMs:
                if str(self.row[0]).split(";")[1] == self.cbOyMs.currentText():
                    copy(str(self.row[0]).split(";")[2])
    def updateOyMTable(self):
        try:
            i = len(self.infoCuadrilla["nombresTecnicos"])
            self.tableTecnicosCuadrilla.setItem(i, 0, QtWidgets.QTableWidgetItem(self.cbOyMs.currentText()))
            self.tableTecnicosCuadrilla.setItem(i, 1, QtWidgets.QTableWidgetItem(self.getTelefonoCorreoyCargoOyM()[0]))
            self.tableTecnicosCuadrilla.setItem(i, 2, QtWidgets.QTableWidgetItem("-"))
            self.tableTecnicosCuadrilla.setItem(i, 3, QtWidgets.QTableWidgetItem(self.getTelefonoCorreoyCargoOyM()[1]))
        except: pass
    def addPersonalTecnicoTable(self):
        try:
            if self.tableTecnicosCuadrilla.rowCount() != 0:
                if self.tableTecnicosCuadrilla.item(self.tableTecnicosCuadrilla.rowCount()-1, 0).text() != "":
                    self.tableTecnicosCuadrilla.setRowCount(self.tableTecnicosCuadrilla.rowCount()+1)
                else: pass
            else:
                if self.cbRBS.currentText() != "":
                    self.tableTecnicosCuadrilla.setRowCount(1)
        except: pass
    def eliminarPersonalTecnicoTable(self):
        try:
            self.tableTecnicosCuadrilla.removeRow(self.tableTecnicosCuadrilla.currentIndex().row())
        except Exception as e:
            pass  
    def eliminarPermisoTable(self):
        try:
            self.tablePermisosNecesarios.removeRow(self.tablePermisosNecesarios.currentIndex().row())
        except Exception as e:
            pass
    def loadcbPermisos(self):
        self.cbPermisos.clear()
        if self.cbRBS.currentText != "":
            self.permisosExistentes = ["", "ECUSITES", "ETC", "Financieros", "PTI", "PRONTO", "SBA", "Telefonica"] # "ATP", "CNT"
            self.cbPermisos.addItems(self.permisosExistentes)
        else: pass
    def loadPermisosTable(self):
        if self.cbPermisos.currentText() != "" and self.cbRBS.currentText() != "":
            self.hasPermisoAlready = 0
            for i in range(self.tablePermisosNecesarios.rowCount()):
                if self.cbPermisos.currentText() == self.tablePermisosNecesarios.item(i, 0).text(): self.hasPermisoAlready = 1
            if not self.hasPermisoAlready:
                self.fecha = datetime.today()
                self.fechaIngreso = str(self.fecha.day)+"/"+str(self.fecha.month)+"/"+str(self.fecha.year)
                self.horaIngreso = "7:00"
                if self.cbPermisos.currentText() == "ECUSITES":
                    self.fecha = datetime.today()+timedelta(days=1)
                else:
                    self.fecha = datetime.today()+timedelta(days=4)
                self.fechaSalida = str(self.fecha.day)+"/"+str(self.fecha.month)+"/"+str(self.fecha.year)
                self.horaSalida = "20:00"
                self.tablePermisosNecesarios.setRowCount(self.tablePermisosNecesarios.rowCount()+1)
                self.tablePermisosNecesarios.setItem(self.tablePermisosNecesarios.rowCount()-1, 0, QtWidgets.QTableWidgetItem(self.cbPermisos.currentText()))
                self.tablePermisosNecesarios.setItem(self.tablePermisosNecesarios.rowCount()-1, 1, QtWidgets.QTableWidgetItem(self.fechaIngreso))
                self.tablePermisosNecesarios.setItem(self.tablePermisosNecesarios.rowCount()-1, 2, QtWidgets.QTableWidgetItem(self.horaIngreso))
                self.tablePermisosNecesarios.setItem(self.tablePermisosNecesarios.rowCount()-1, 3, QtWidgets.QTableWidgetItem(self.fechaSalida))
                self.tablePermisosNecesarios.setItem(self.tablePermisosNecesarios.rowCount()-1, 4, QtWidgets.QTableWidgetItem(self.horaSalida))
                self.cbPermisos.setCurrentIndex(0)
    def clearDataFromRBS(self):
        self.cbOyMs.clear()
        self.cbCuadrillas.clear()
        self.cbCuadrillas.clear()
        self.tablePermisosNecesarios.clear()
        self.tableTecnicosCuadrilla.clear()
        self.lineRegion.setText("")
        self.lineZona.setText("")
        self.lineProvincia.setText("")
    def loadDataFromRBS(self):
        self.clearDataFromRBS()
        self.cuadrillas = [""]
        self.OyMs = [""]
        bandera = 0
        if self.cbRBS.currentIndex() != 0:
            with open("data/RBS.csv") as self.csvRBS:
                self.InfoRBS = reader(self.csvRBS)
                for self.row in self.InfoRBS:
                    # RBS Found
                    try:
                        if str(self.row[0]).split(";")[1] == self.cbRBS.currentText():
                            self.OyMs[0] = str(self.row[0]).split(";")[6]
                            self.cuadrillas[0] = str(self.row[0]).split(";")[7]     
                            self.lineRegion.setText("R" + str(self.getRegionZonaRBS()[0]))
                            self.lineZona.setText("Z" + str(self.getRegionZonaRBS()[1]))
                            self.lineProvincia.setText(self.getProvinciaCantonRBS()[0])
                        else: pass
                    except: 
                        bandera +=1
                        print(bandera)
                        pass
                    # Collect all Cuadrillas
                    try:
                        if (str(self.row[0]).split(";")[7]) not in self.cuadrillas and str(self.row[0]).split(";")[7] != "" and str(self.row[0]).split(";")[7] != "Nan" and str(self.row[0]).split(";")[7].upper() !="Cuadrilla".upper():
                            self.cuadrillas.append(str(self.row[0]).split(";")[7])
                        else: pass
                    except:pass
                    # Collect all OyMs
                    try:
                        if (str(self.row[0]).split(";")[6]) not in self.OyMs and str(self.row[0]).split(";")[6] != "" and str(self.row[0]).split(";")[6].upper() !="O&M Dueño".upper():
                            self.OyMs.append(str(self.row[0]).split(";")[6])
                        else: pass
                    except:pass
            # Load in ComboBox Cuadrillas and O&Ms
            for i in self.cuadrillas:
                self.cbCuadrillas.addItem(QtGui.QIcon("ResourcesFolder/featherIcons/users.svg"), i)
            for i in self.OyMs:
                self.cbOyMs.addItem(QtGui.QIcon("ResourcesFolder/featherIcons/user.svg"), i)
            self.loadSupervisorNTableCuadrilla()
        else: self.loadcbPermisos()
        self.tableTecnicosCuadrilla.setHorizontalHeaderLabels(["Nombre", "Teléfono", "Cédula", "Correo"])
        self.tablePermisosNecesarios.setHorizontalHeaderLabels(["Permiso", "Fecha Ingreso", "Hora Ingreso", "Fecha Salida", "Hora Salida"])
    def loadSupervisorNTableCuadrilla(self):
        self.cbSupervisores.clear()
        self.tableTecnicosCuadrilla.clear()
        self.supervisores = [""]
        self.infoCuadrilla = {}
        self.cargoTecnicos = []
        self.nombresTecnicos = []
        self.cedulaTecnicos = []
        self.telefonoTecnicos = []
        self.correoTecnicos = []
        self.placaTecnicos = []
        self.llaveTecnicos = []
        with open("data/Cuadrillas.csv") as self.csvCuadrillas:
            self.InfoCuadrillas = reader(self.csvCuadrillas)
            for self.row in self.InfoCuadrillas:
                if str(self.row[0]).split(";")[2] == self.cbCuadrillas.currentText():
                    if str(self.row[0]).split(";")[1] != "":
                        self.supervisores[0] = str(self.row[0]).split(";")[1]
                        self.cargoTecnicos.append(str(self.row[0]).split(";")[3])
                        self.nombresTecnicos.append(str(self.row[0]).split(";")[4])
                        self.cedulaTecnicos.append(str(self.row[0]).split(";")[5])
                        self.telefonoTecnicos.append(str(self.row[0]).split(";")[6])
                        self.correoTecnicos.append(str(self.row[0]).split(";")[7])
                        self.placaTecnicos.append(str(self.row[0]).split(";")[8])
                        self.llaveTecnicos.append(str(self.row[0]).split(";")[9])
                    else: pass
                else:pass
                if (str(self.row[0]).split(";")[1]) not in self.supervisores and str(self.row[0]).split(";")[1].upper() != "WIRELESS" and str(self.row[0]).split(";")[1].upper() != "ALU" and str(self.row[0]).split(";")[1].upper() != "GENERADORES" and str(self.row[0]).split(";")[1] != "" and str(self.row[0]).split(";")[1].upper() !="Supervisor".upper():
                    self.supervisores.append(str(self.row[0]).split(";")[1])
                else: pass
        for i in self.supervisores:
                self.cbSupervisores.addItem(QtGui.QIcon("ResourcesFolder/featherIcons/user.svg"), i)
        self.infoCuadrilla["supervisor"] = self.cbSupervisores.currentText()
        self.infoCuadrilla["cuadrilla"] = self.cbCuadrillas.currentText()
        self.infoCuadrilla["cargoTecnicos"] = self.cargoTecnicos
        self.infoCuadrilla["nombresTecnicos"] = self.nombresTecnicos
        self.infoCuadrilla["cedulaTecnicos"] = self.cedulaTecnicos
        self.infoCuadrilla["telefonoTecnicos"] = self.telefonoTecnicos
        self.infoCuadrilla["correoTecnicos"] = self.correoTecnicos
        self.infoCuadrilla["placasTecnicos"] = self.placaTecnicos
        self.infoCuadrilla["llaveTecnicos"] = self.llaveTecnicos
        # Loading Table Tecnicos
        self.tableTecnicosCuadrilla.setRowCount(len(self.infoCuadrilla["nombresTecnicos"])+1)
        for i in range(len(self.infoCuadrilla["nombresTecnicos"])):
            self.tableTecnicosCuadrilla.setItem(i, 0, QtWidgets.QTableWidgetItem(self.infoCuadrilla["nombresTecnicos"][i]))
            self.tableTecnicosCuadrilla.setItem(i, 1, QtWidgets.QTableWidgetItem(self.infoCuadrilla["telefonoTecnicos"][i]))
            self.tableTecnicosCuadrilla.setItem(i, 2, QtWidgets.QTableWidgetItem(self.infoCuadrilla["cedulaTecnicos"][i]))
            self.tableTecnicosCuadrilla.setItem(i, 3, QtWidgets.QTableWidgetItem(self.infoCuadrilla["correoTecnicos"][i]))
        try:
            self.tableTecnicosCuadrilla.setItem(i+1, 0, QtWidgets.QTableWidgetItem(self.cbOyMs.currentText()))
            self.tableTecnicosCuadrilla.setItem(i+1, 1, QtWidgets.QTableWidgetItem(self.getTelefonoCorreoyCargoOyM()[0]))
            self.tableTecnicosCuadrilla.setItem(i+1, 2, QtWidgets.QTableWidgetItem("-"))
            self.tableTecnicosCuadrilla.setItem(i+1, 3, QtWidgets.QTableWidgetItem(self.getTelefonoCorreoyCargoOyM()[1]))
        except: pass
        # Load Table Permisos
        self.tablePermisosNecesarios.setRowCount(0)
        self.fixHeaderTables()
    def fixHeaderTables(self):
        self.tableTecnicosCuadrilla.setHorizontalHeaderLabels(["Nombre", "Teléfono", "Cédula", "Correo"])
        self.tablePermisosNecesarios.setHorizontalHeaderLabels(["Permiso", "Fecha Ingreso", "Hora Ingreso", "Fecha Salida", "Hora Salida"])
        #Font
        afont = QtGui.QFont()
        afont.setFamily("Arial")
        afont.setPointSize(13)
        self.tablePermisosNecesarios.horizontalHeaderItem(0).setFont(afont)
        self.tablePermisosNecesarios.horizontalHeaderItem(1).setFont(afont)
        self.tablePermisosNecesarios.horizontalHeaderItem(2).setFont(afont)
        self.tablePermisosNecesarios.horizontalHeaderItem(3).setFont(afont)
        self.tablePermisosNecesarios.horizontalHeaderItem(4).setFont(afont)
        self.tableTecnicosCuadrilla.horizontalHeaderItem(0).setFont(afont)
        self.tableTecnicosCuadrilla.horizontalHeaderItem(1).setFont(afont)
        self.tableTecnicosCuadrilla.horizontalHeaderItem(2).setFont(afont)
        self.tableTecnicosCuadrilla.horizontalHeaderItem(3).setFont(afont)
        # Icon
        # Permisos
        self.tablePermisosNecesarios.horizontalHeaderItem(0).setIcon(QtGui.QIcon("ResourcesFolder/featherIcons/file-text.svg"))
        self.tablePermisosNecesarios.horizontalHeaderItem(1).setIcon(QtGui.QIcon("ResourcesFolder/featherIcons/calendar.svg"))
        self.tablePermisosNecesarios.horizontalHeaderItem(2).setIcon(QtGui.QIcon("ResourcesFolder/featherIcons/clock.svg"))
        self.tablePermisosNecesarios.horizontalHeaderItem(3).setIcon(QtGui.QIcon("ResourcesFolder/featherIcons/calendar.svg"))
        self.tablePermisosNecesarios.horizontalHeaderItem(4).setIcon(QtGui.QIcon("ResourcesFolder/featherIcons/clock.svg"))
        # Tecnicos
        self.tableTecnicosCuadrilla.horizontalHeaderItem(0).setIcon(QtGui.QIcon("ResourcesFolder/featherIcons/user.svg"))
        self.tableTecnicosCuadrilla.horizontalHeaderItem(1).setIcon(QtGui.QIcon("ResourcesFolder/featherIcons/unlock.svg"))
        self.tableTecnicosCuadrilla.horizontalHeaderItem(2).setIcon(QtGui.QIcon("ResourcesFolder/featherIcons/briefcase.svg"))
        self.tableTecnicosCuadrilla.horizontalHeaderItem(3).setIcon(QtGui.QIcon("ResourcesFolder/featherIcons/at-sign.svg"))
    def adjustTablesSize(self):
        # Table Tecnicos
        self.tableTecnicosCuadrilla.setColumnWidth(0, int(self.tableTecnicosCuadrilla.width()*34.5/100)) #34.5%
        self.tableTecnicosCuadrilla.setColumnWidth(1, int(self.tableTecnicosCuadrilla.width()*15.5/100)) #15.5%
        self.tableTecnicosCuadrilla.setColumnWidth(2, int(self.tableTecnicosCuadrilla.width()*15.2/100)) #15.5%
        self.tableTecnicosCuadrilla.setColumnWidth(3, int(self.tableTecnicosCuadrilla.width()*34.5/100)) #34.5%
        self.tableTecnicosCuadrilla.setItemDelegateForColumn(0, AlignDelegate(self.tableTecnicosCuadrilla))
        self.tableTecnicosCuadrilla.setItemDelegateForColumn(1, AlignDelegate(self.tableTecnicosCuadrilla))
        self.tableTecnicosCuadrilla.setItemDelegateForColumn(2, AlignDelegate(self.tableTecnicosCuadrilla))
        self.tableTecnicosCuadrilla.setItemDelegateForColumn(3, AlignDelegate(self.tableTecnicosCuadrilla))
        # Table Permisos
        self.tablePermisosNecesarios.setColumnWidth(0, int(self.tablePermisosNecesarios.width()*20/100)) #20%
        self.tablePermisosNecesarios.setColumnWidth(1, int(self.tablePermisosNecesarios.width()*20/100)) #20%
        self.tablePermisosNecesarios.setColumnWidth(2, int(self.tablePermisosNecesarios.width()*20/100)) #20%
        self.tablePermisosNecesarios.setColumnWidth(3, int(self.tablePermisosNecesarios.width()*20/100)) #20%
        self.tablePermisosNecesarios.setColumnWidth(4, int(self.tablePermisosNecesarios.width()*20/100)) #20%
        self.tablePermisosNecesarios.setItemDelegateForColumn(0, ReadOnlyDelegate(self.tablePermisosNecesarios))
        self.tablePermisosNecesarios.setItemDelegateForColumn(0, AlignDelegate(self.tablePermisosNecesarios))
        self.tablePermisosNecesarios.setItemDelegateForColumn(1, AlignDelegate(self.tablePermisosNecesarios))
        self.tablePermisosNecesarios.setItemDelegateForColumn(2, AlignDelegate(self.tablePermisosNecesarios))
        self.tablePermisosNecesarios.setItemDelegateForColumn(3, AlignDelegate(self.tablePermisosNecesarios))
        self.tablePermisosNecesarios.setItemDelegateForColumn(4, AlignDelegate(self.tablePermisosNecesarios))        
    def loadRBS(self):
        self.cbRBS.clear()
        with open("data/RBS.csv") as self.csvRBS:
            self.InfoRBS = reader(self.csvRBS)
            for self.row in self.InfoRBS:
                if str(self.row[0]).split(";")[1] != "":
                    if str(self.row[0]).split(";")[1] =="RBS":
                        self.cbRBS.addItem(QtGui.QIcon(), "")
                    else:    
                        self.cbRBS.addItem(QtGui.QIcon("ResourcesFolder/featherIcons/briefcase.svg"), str(self.row[0]).split(";")[1])
        pass
    def maximizeWindow(self):
        if not self.isMaximized():self.showMaximized()
        else: self.showNormal()
        self.adjustTablesSize()
    def closeWindow(self):
        deleteAllData()
        self.close()
    def mouseMoveEvent(self, event):
        try:
            if self.oldPosition.y()-self.y() <= 45 and not self.isMaximized():
                p = event.globalPosition()
                delta = QtCore.QPoint(p.toPoint()-self.oldPosition)
                self.move(self.x()+delta.x(), self.y()+delta.y())
                p = event.globalPosition()
                self.oldPosition = p.toPoint()
        except Exception as e:
            pass
    def mousePressEvent(self, event):
        try:
            p = event.globalPosition()
            self.oldPosition = p.toPoint()
        except Exception as e:
            pass
    def verificationInformation(self):
        self.error = []
        # Verification of ComboBoxes
        if self.cbRBS.currentText() == "": self.error.append(1) # Sin seleccion de RBS
        else: pass
        if self.cbOyMs.currentText() == "": self.error.append(2) # Sin seleccion de OYM
        else: pass
        if self.cbSupervisores.currentText() == "": self.error.append(3) # Sin seleccion de supervisores
        else: pass
        if self.cbCuadrillas.currentText() == "": self.error.append(4) # Sin seleccion de cuadrilla
        else: pass
        # Verificacion de Alarma
        if self.lineAlarma.text().strip() == "": self.error.append(5) # Sin Alarma
        else: pass
        # Verificacion de Tablas
        # Verificacion Tabla Permisos
        if self.tablePermisosNecesarios.rowCount() == 0: self.error.append(6) # Sin Permisos
        else:
            for i in range(self.tablePermisosNecesarios.rowCount()):
                if not self.fechaValida(self.tablePermisosNecesarios.item(i, 1).text()): self.error.append(6) # Fecha Ingreso Invalida
                else: pass
                if not self.horaValida(self.tablePermisosNecesarios.item(i, 2).text()): self.error.append(7) # Hora Ingreso Invalida
                else: pass
                if not self.fechaValida(self.tablePermisosNecesarios.item(i, 3).text()): self.error.append(8) # Fecha Salida Invalida
                else: pass
                if not self.horaValida(self.tablePermisosNecesarios.item(i, 4).text()): self.error.append(9) # Hora Salida Invalida
                else: pass
        # Verificacion Tabla Tecnicos
        if self.tableTecnicosCuadrilla.rowCount() == 0: self.error.append(10) # Sin Permisos
        else:
            for i in range(self.tableTecnicosCuadrilla.rowCount()):
                if not self.nombreValido(self.tableTecnicosCuadrilla.item(i, 0).text()): self.error.append(11) # Nombre Inválido
                else: pass
                if not self.telefonoValido(self.tableTecnicosCuadrilla.item(i, 1).text()): self.error.append(12) # Telefono Inválido
                else: pass
                if not self.cedulaValido(self.tableTecnicosCuadrilla.item(i, 2).text()): self.error.append(13) # Cedula Inválido
                else: pass
                if not self.correoValido(self.tableTecnicosCuadrilla.item(i, 3).text()): self.error.append(14) # Cedula Inválido
                else: pass
    def telefonoValido(self, telefono:str)->bool:
        if telefono.isnumeric():
            if len(telefono) == 10 or len(telefono) == 9:
                return True
        elif telefono == "-" or telefono.strip() == "":return True
        else: return False
    def cedulaValido(self, cedula:str)->bool:
        if len(cedula) == 9 or len(cedula) == 10 or cedula == "-" or cedula.strip() == "":
            if cedula == "-" or cedula.strip() == "":
                return True
            elif cedula.isnumeric():
                return True
            else: return False
        else: return False
    def correoValido(self, correo:str)->bool:
        if len(correo.split("@")) == 2:
            if len(correo.split("@")[0]) > 0:
                if len(correo.split("@")[1]) > 0:
                    if len(correo.split("@")[1].split(".")) >= 2:
                        if len(correo.split("@")[1].split(".")[0]) > 0:
                            if len(correo.split("@")[1].split(".")[1]) > 0:
                                return True
                            else: pass
                        else: pass
                    else: pass
                else: pass
            else: pass
        else: pass
        if correo == "-" or correo.strip() == "":return True
        else: return False
    def nombreValido(self, nombre:str)->bool:
        if nombre.replace(" ", "").isalpha(): return True
        return False
    def horaValida(self, hora:str)->bool:
        if len(hora) >= 4:
            if len(hora.split(":")) == 2:
                if len(hora.split(":")[0]) >= 1:
                    if len(hora.split(":")[1]) >= 1:
                        if hora.split(":")[0].isnumeric() and hora.split(":")[1].isnumeric():
                            if int(hora.split(":")[0]) >= 0 and int(hora.split(":")[0]) < 24:
                                if int(hora.split(":")[1]) >= 0 and int(hora.split(":")[1]) < 60:
                                    return True
                                else: return False
                            else: return False
                        else: return False
                    else: return False
                else: return False
            else: return False
        else: return False
    def fechaValida(self, fecha:str)->bool:
        if len(fecha.split("/")) == 3 :
            if fecha.split("/")[0].isnumeric() and fecha.split("/")[1].isnumeric() and fecha.split("/")[2].isnumeric():
                try:
                    datetime(year=int(fecha.split("/")[2]), month=int(fecha.split("/")[1]), day=int(fecha.split("/")[0]))
                    return True
                except:
                    return False
            else: return False
        else: return False
    def replace_text_in_paragraph(self, paragraph, key, value):
        if key in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                #print(item.text)
                if key in item.text:
                    item.text = item.text.replace(key, value)
    def getProvinciaCantonRBS(self):
        with open("data/RBS.csv") as self.csvRBS:
            self.InfoRBS = reader(self.csvRBS)
            for self.row in self.InfoRBS:
                if str(self.row[0]).split(";")[1] == self.cbRBS.currentText():
                    return [str(self.row[0]).split(";")[4], str(self.row[0]).split(";")[5]]
    def getCargoTecnico(self, nombreTecnico:str):
        if nombreTecnico == self.cbOyMs.currentText():
            return self.getTelefonoCorreoyCargoOyM()[2]
        elif nombreTecnico == "":
            return ""
        else:
            with open("data/Cuadrillas.csv") as self.csvRBS:
                self.InfoRBS = reader(self.csvRBS)
                for self.row in self.InfoRBS:
                    if str(self.row[0]).split(";")[4] == nombreTecnico:
                        return str(self.row[0]).split(";")[3]
            return ""
    def getTelefonoCorreoyCargoOyM(self):
        with open("data/OyMs.csv") as self.csvOYMs:
            self.InfoOYMs = reader(self.csvOYMs)
            for self.row in self.InfoOYMs:
                if str(self.row[0]).split(";")[1] == self.cbOyMs.currentText():
                    return [str(self.row[0]).split(";")[2], str(self.row[0]).split(";")[3], str(self.row[0]).split(";")[4]]
        return ["-", "-", "-"]
    def getTelefonoyCorreoSupervisor(self):
        with open("data/Supervisores.csv") as self.csvSupervisores:
            self.InfoSupervisores = reader(self.csvSupervisores)
            for self.row in self.InfoSupervisores:
                if str(self.row[0]).split(";")[1] == self.cbSupervisores.currentText():
                    return [str(self.row[0]).split(";")[2], str(self.row[0]).split(";")[3]]
        return ["", ""]
    def makePermisoFinanciero(self, nombre:str):
        book = load_workbook("Formatos/PERMISOS FINANCIEROS.xlsx")
        sheet = book.active
        sheet["E9"].value = self.getProvinciaCantonRBS()[0] # Provincia
        sheet["E10"].value = self.permisos[nombre][0] # Fecha Ingreso
        sheet["E12"].value = self.cbRBS.currentText() # nameRBS
        sheet["E13"].value = self.permisos[nombre][0] # Fecha Ingreso
        sheet["E14"].value = self.permisos[nombre][2] # Fecha Salida
        sheet["E15"].value = self.permisos[nombre][1] + " - " + self.permisos[nombre][3] # Hora Entrada - Salida
        sheet["E16"].value = self.lineAlarma.text() # Alarma
        sheet["D20"].value = self.tableTecnicosCuadrilla.item(0, 0).text() # team leader nombre
        sheet["E20"].value = self.tableTecnicosCuadrilla.item(0, 2).text() # Team Leader CC
        sheet["F20"].value = "RHELEC"
        for i in range(self.tableTecnicosCuadrilla.rowCount()-1):
            sheet["D"+str(21+i)].value = self.tableTecnicosCuadrilla.item(i+1, 0).text() # Nombres Tecnicos
            sheet["E"+str(21+i)].value = self.tableTecnicosCuadrilla.item(i+1, 2).text() # CC Tecnicos
            if self.tableTecnicosCuadrilla.item(i+1, 0).text() == self.cbOyMs.currentText():
                sheet["F"+str(21+i)].value = "Claro"
            else:
                sheet["F"+str(21+i)].value = "RHELEC"
        pathOutputFile = "Output/"
        nameOutputFile = "PERMISOS FINANCIEROS - " + self.cbRBS.currentText() + ".xlsx"
        book.save(pathOutputFile+nameOutputFile)
        book.close
        return pathOutputFile + nameOutputFile
    def makePermisoSBA(self):
        book = load_workbook("Formatos/FORMATO SITIOS SBA RBS.xlsx")
        sheet = book.active
        sheet["D6"].value = self.cbRBS.currentText() # NameRBS
        sheet["E3"].value = self.permisos["SBA"][0] # Fecha Desde
        sheet["D11"].value = self.permisos["SBA"][0] # Fecha Desde
        sheet["D12"].value = self.permisos["SBA"][2] # Fecha Salida
        sheet["C27"].value = self.lineAlarma.text() # Actividad
        sheet["D13"].value = self.tableTecnicosCuadrilla.item(0, 0).text() # Nombre Team Leader
        sheet["D14"].value = self.tableTecnicosCuadrilla.item(0, 1).text() # Telefono Team Leader
        for i in range(self.tableTecnicosCuadrilla.rowCount()):
            sheet["C"+str(19+i)].value = self.tableTecnicosCuadrilla.item(i, 0).text() # nombre Tecnico
            sheet["D"+str(19+i)].value = self.tableTecnicosCuadrilla.item(i, 2).text() # cc tecnico
            sheet["E"+str(19+i)].value = self.tableTecnicosCuadrilla.item(i, 1).text() # telefono tecnico
            if self.tableTecnicosCuadrilla.item(i, 0).text() == self.cbOyMs.currentText():
                sheet["F"+str(19+i)].value = "Claro"
            else:
                sheet["F"+str(19+i)].value = "RHELEC"
        pathOutputFile = "Output/"
        nameOutputFile = "FORMATO SITIOS SBA RBS " + self.cbRBS.currentText() + ".xlsx"
        book.save(pathOutputFile+nameOutputFile)
        book.close
        return pathOutputFile + nameOutputFile
    def makePermisoETC(self):
        self.template_file_path = 'Formatos/FORMATO SITIOS ETC.docx'
        self.auxNombreTecnicos = ["", "", "", "", ""]
        self.auxCCTecnicos = ["", "", "", "", ""]
        self.auxTelefonoTecnicos = ["", "", "", "", ""]
        self.auxEmailTecnicos = ["", "", "", "", ""]
        self.auxEmpresaTecnicos = ["", "", "", "", ""]
        for i in range(self.tableTecnicosCuadrilla.rowCount()):
            self.auxNombreTecnicos[i] = self.tableTecnicosCuadrilla.item(i, 0).text()
            self.auxCCTecnicos[i] = self.tableTecnicosCuadrilla.item(i, 2).text()
            self.auxTelefonoTecnicos[i] = self.tableTecnicosCuadrilla.item(i, 1).text()
            self.auxEmailTecnicos[i] = self.tableTecnicosCuadrilla.item(i, 3).text()
            if self.tableTecnicosCuadrilla.item(i, 0).text() == self.cbOyMs.currentText():
                self.auxEmpresaTecnicos[i] = "Claro"
            else:
                self.auxEmpresaTecnicos[i] = "RHELEC"
        self.variables = {"fechayHoraIngreso": self.permisos["ETC"][0] + " " + self.permisos["ETC"][1],
                    "fechayHoraSalida": self.permisos["ETC"][2] + " " + self.permisos["ETC"][3],
                    "nameRBS": self.cbRBS.currentText(),
                    "Alarma": self.lineAlarma.text(),
                    "nombreTecnico1": self.auxNombreTecnicos[0],
                    "nombreTecnico2": self.auxNombreTecnicos[1],
                    "nombreTecnico3": self.auxNombreTecnicos[2],
                    "nombreTecnico4": self.auxNombreTecnicos[3],
                    "nombreTecnico5": self.auxNombreTecnicos[4],
                    "CCTec1": self.auxCCTecnicos[0],
                    "CCTec2": self.auxCCTecnicos[1],
                    "CCTec3": self.auxCCTecnicos[2],
                    "CCTec4": self.auxCCTecnicos[3],
                    "CCTec5": self.auxCCTecnicos[4],
                    "TelTec1": self.auxTelefonoTecnicos[0],
                    "TelTec2": self.auxTelefonoTecnicos[1],
                    "TelTec3": self.auxTelefonoTecnicos[2],
                    "TelTec4": self.auxTelefonoTecnicos[3],
                    "TelTec5": self.auxTelefonoTecnicos[4],
                    "emailTec1": self.auxEmailTecnicos[0],
                    "emailTec2": self.auxEmailTecnicos[1],
                    "emailTec3": self.auxEmailTecnicos[2],
                    "emailTec4": self.auxEmailTecnicos[3],
                    "emailTec5": self.auxEmailTecnicos[4],
                    "Cat1": self.getCargoTecnico(self.auxNombreTecnicos[0]),
                    "Cat2": self.getCargoTecnico(self.auxNombreTecnicos[1]),
                    "Cat3": self.getCargoTecnico(self.auxNombreTecnicos[2]),
                    "Cat4": self.getCargoTecnico(self.auxNombreTecnicos[3]),
                    "Cat5": self.getCargoTecnico(self.auxNombreTecnicos[4]),
                    "Emt1": self.auxEmpresaTecnicos[0],
                    "Emt2": self.auxEmpresaTecnicos[1],
                    "Emt3": self.auxEmpresaTecnicos[2],
                    "Emt4": self.auxEmpresaTecnicos[3],
                    "Emt5": self.auxEmpresaTecnicos[4]}
        self.template_document = Document(self.template_file_path)
        self.changeVariablesInWord()
        output_file_path = 'Output/FORMATO SITIOS ETC - ' + self.cbRBS.currentText() + '.docx'
        self.template_document.save(output_file_path)
        return output_file_path
    def makePermisoATB(self):
        return True
    def makePermisoECUSITES(self):
        book = load_workbook("Formatos/SITES - FORMULARIO DE INGRESO.xlsx")
        sheet = book.active
        sheet["D5"].value = self.getProvinciaCantonRBS()[0] # Provincia RBS
        sheet["D6"].value = self.cbRBS.currentText() # name RBS
        sheet["D7"].value = self.getProvinciaCantonRBS()[0] # Provincia RBS
        sheet["L5"].value = self.permisos["ECUSITES"][0] # Fecha Ingreso
        sheet["I39"].value = self.permisos["ECUSITES"][0]+" "+self.permisos["ECUSITES"][1] # Fecha y hora Ingreso
        sheet["L39"].value = self.permisos["ECUSITES"][2]+" "+self.permisos["ECUSITES"][3] # Fecha y hora Salida
        for i in range(self.tableTecnicosCuadrilla.rowCount()):
            sheet["D"+str(10+i)].value = self.tableTecnicosCuadrilla.item(i, 0).text() # Nombre Tecnicos
            sheet["K"+str(10+i)].value = self.tableTecnicosCuadrilla.item(i, 1).text() # telefono tecnico
            sheet["D"+str(19+i)].value = self.tableTecnicosCuadrilla.item(i, 0).text() # Nombre Tecnicos
            sheet["K"+str(19+i)].value = self.tableTecnicosCuadrilla.item(i, 1).text() # telefono tecnico
        pathOutputFile = "Output/"
        nameOutputFile = "SITES - FORMULARIO DE INGRESO - " + self.cbRBS.currentText() + ".xlsx"
        book.save(pathOutputFile+nameOutputFile)
        book.close
        return pathOutputFile + nameOutputFile
    def makePermisoCNT(self):
        self.template_file_path = 'Formatos/FORMULARIO INGRESO CNT.docx'
        self.auxNombreTecnicos = ["", "", "", "", ""]
        self.auxCCTecnicos = ["", "", "", "", ""]
        self.auxTelefonoTecnicos = ["", "", "", "", ""]
        self.auxEmailTecnicos = ["", "", "", "", ""]
        self.auxEmpresaTecnicos = ["", "", "", "", ""]
        for i in range(self.tableTecnicosCuadrilla.rowCount()):
            self.auxNombreTecnicos[i] = self.tableTecnicosCuadrilla.item(i, 0).text()
            self.auxCCTecnicos[i] = self.tableTecnicosCuadrilla.item(i, 2).text()
            self.auxTelefonoTecnicos[i] = self.tableTecnicosCuadrilla.item(i, 1).text()
            self.auxEmailTecnicos[i] = self.tableTecnicosCuadrilla.item(i, 3).text()
            if self.tableTecnicosCuadrilla.item(i, 0).text() == self.cbOyMs.currentText():
                self.auxEmpresaTecnicos[i] = "Claro"
            else:
                self.auxEmpresaTecnicos[i] = "RHELEC"
        self.variables = {"fechaIng": self.permisos["CNT"][0] + " " + self.permisos["CNT"][1],
                    "fechaSal": self.permisos["CNT"][2] + " " + self.permisos["CNT"][3],
                    "nombreOYM": self.cbOyMs.currentText(),
                    "fechaRetiroLlaves": str(datetime(day=int(self.permisos["CNT"][0].split("/")[0]), month=int(self.permisos["CNT"][0].split("/")[1]), year=int(self.permisos["CNT"][0].split("/")[2]))-timedelta(2)),
                    "fechaEntregaLlaves": self.permisos["CNT"][2],
                    "telefonoOYM":self.getTelefonoCorreoyCargoOyM()[0],
                    "nameRBSFija": self.cbRBS.currentText(),
                    "AlarmaRBS": self.lineAlarma.text(),
                    "provinciaRBSFija":self.getProvinciaCantonRBS()[0],
                    "nombreTec1": self.auxNombreTecnicos[0],
                    "nombreTec2": self.auxNombreTecnicos[1],
                    "nombreTec3": self.auxNombreTecnicos[2],
                    "nombreTec4": self.auxNombreTecnicos[3],
                    "nombreTec5": self.auxNombreTecnicos[4],
                    "ccTec1": self.auxCCTecnicos[0],
                    "ccTec2": self.auxCCTecnicos[1],
                    "ccTec3": self.auxCCTecnicos[2],
                    "ccTec4": self.auxCCTecnicos[3],
                    "ccTec5": self.auxCCTecnicos[4],
                    "telTec1": self.auxTelefonoTecnicos[0],
                    "telTec2": self.auxTelefonoTecnicos[1],
                    "telTec3": self.auxTelefonoTecnicos[2],
                    "telTec4": self.auxTelefonoTecnicos[3],
                    "telTec5": self.auxTelefonoTecnicos[4],
                    "empTec1": self.auxEmpresaTecnicos[0],
                    "empTec2": self.auxEmpresaTecnicos[1],
                    "empTec3": self.auxEmpresaTecnicos[2],
                    "empTec4": self.auxEmpresaTecnicos[3],
                    "empTec5": self.auxEmpresaTecnicos[4]}
        self.template_document = Document(self.template_file_path)
        self.changeVariablesInWord()
        output_file_path = 'Output/FORMULARIO INGRESO CNT - ' + self.cbRBS.currentText() + '.docx'
        self.template_document.save(output_file_path)
        return output_file_path
        pass
    def makePermisoTELEFONICA(self):
        self.template_file_path = 'Formatos/PERMISOS MOVISTAR.docx'
        self.auxNombreTecnicos = ["", "", "", "", ""]
        self.auxCCTecnicos = ["", "", "", "", ""]
        self.auxTelefonoTecnicos = ["", "", "", "", ""]
        for i in range(self.tableTecnicosCuadrilla.rowCount()):
            self.auxNombreTecnicos[i] = self.tableTecnicosCuadrilla.item(i, 0).text()
            self.auxCCTecnicos[i] = self.tableTecnicosCuadrilla.item(i, 2).text()
            self.auxTelefonoTecnicos[i] = self.tableTecnicosCuadrilla.item(i, 1).text()
        self.variables = {"${fechayHoraIngreso}": self.permisos["Telefonica"][0] + " " + self.permisos["Telefonica"][1],
                    "${nombreTecnico1}": self.auxNombreTecnicos[0],
                    "${nombreTecnico2}": self.auxNombreTecnicos[1],
                    "${nombreTecnico3}": self.auxNombreTecnicos[2],
                    "${nombreTecnico4}": self.auxNombreTecnicos[3],
                    "${nombreTecnico5}": self.auxNombreTecnicos[4],
                    "${ccTecnico1}": self.auxCCTecnicos[0],
                    "${ccTecnico2}": self.auxCCTecnicos[1],
                    "${ccTecnico3}": self.auxCCTecnicos[2],
                    "${ccTecnico4}": self.auxCCTecnicos[3],
                    "${ccTecnico5}": self.auxCCTecnicos[4],
                    "${telefonoTecnico1}": self.auxTelefonoTecnicos[0],
                    "${telefonoTecnico2}": self.auxTelefonoTecnicos[1],
                    "${telefonoTecnico3}": self.auxTelefonoTecnicos[2],
                    "${telefonoTecnico4}": self.auxTelefonoTecnicos[3],
                    "${telefonoTecnico5}": self.auxTelefonoTecnicos[4],
                    "nombredelaradiobase": self.cbRBS.currentText(),
                    "${Alarma}": self.lineAlarma.text(),
                    "${fechaIngreso}": self.permisos["Telefonica"][0],
                    "${fechaSalida}": self.permisos["Telefonica"][2],
                    "${nombreOyM}": self.cbOyMs.currentText(),
                    "${telefonoOyM}": self.getTelefonoCorreoyCargoOyM()[0]}
        self.template_document = Document(self.template_file_path)
        self.changeVariablesInWord()
        '''for i in variables:
            for p in template_document.paragraphs:
                if p.text.find(i)>=0:
                    p.text=p.text.replace(i,variables[i])'''
        '''for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                self.replace_text_in_paragraph(paragraph, variable_key, variable_value)
            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            #print(paragraph.text)
                            self.replace_text_in_paragraph(paragraph, variable_key, variable_value)'''
        output_file_path = 'Output/PERMISOS MOVISTAR - ' + self.cbRBS.currentText() + '.docx'
        self.template_document.save(output_file_path)
        return output_file_path
    def changeVariablesInWord(self):
        for i in self.variables:
            for p in self.template_document.paragraphs:
                if p.text.find(i)>=0:
                    p.text=p.text.replace(i,self.variables[i])
        for variable_key, variable_value in self.variables.items():
            for paragraph in self.template_document.paragraphs:
                self.replace_text_in_paragraph(paragraph, variable_key, variable_value)
            for table in self.template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            #print(paragraph.text)
                            self.replace_text_in_paragraph(paragraph, variable_key, variable_value)
    def makePermisoPTI(self):
        book = load_workbook("Formatos/SOLICITUD INGRESO PTI.xlsx")
        sheet = book.active
        sheet["D7"].value = self.permisos["PTI"][0] # Fecha Ingreso
        sheet["C9"].value = self.cbRBS.currentText() # name RBS
        sheet["E9"].value = self.getProvinciaCantonRBS()[0] # provincia RBS
        sheet["F9"].value = self.getProvinciaCantonRBS()[1] # canton RBS
        sheet["G9"].value = self.lineAlarma.text() # Alarma
        sheet["G17"].value = self.permisos["PTI"][1] # Hora Ingreso
        sheet["H17"].value = self.permisos["PTI"][3] # Hora Salida
        sheet["C23"].value = self.lineAlarma.text() # Alarma
        sheet["F22"].value = "" # Placa
        sheet["H22"].value = self.permisos["PTI"][0] # Fecha Ingreso
        sheet["H23"].value = self.permisos["PTI"][2] # Fecha Salida
        for i in range(self.tableTecnicosCuadrilla.rowCount()):
            sheet["B"+str(27+i)].value = self.tableTecnicosCuadrilla.item(i, 0).text() # Nombre Tecnicos
            sheet["C"+str(27+i)].value = self.tableTecnicosCuadrilla.item(i, 2).text() # CC Tecnicos
            if self.tableTecnicosCuadrilla.item(i, 0).text() == self.cbOyMs.currentText():
                sheet["F"+str(27+i)].value = "Claro" # Empres
            else:
                sheet["F"+str(27+i)].value = "RHELEC" # Empres
            sheet["H"+str(27+i)].value = self.tableTecnicosCuadrilla.item(i, 1).text() # Telefono Tecnicos
        sheet["C39"].value = self.tableTecnicosCuadrilla.item(0, 0).text() # Nombre Team Leader
        sheet["E39"].value = self.tableTecnicosCuadrilla.item(0, 2).text() # CC team leader
        sheet["G39"].value = self.tableTecnicosCuadrilla.item(0, 3).text() # correo team leader
        sheet["C53"].value = self.tableTecnicosCuadrilla.item(0, 0).text() # nombre team Leader
        sheet["C54"].value = self.getCargoTecnico(self.tableTecnicosCuadrilla.item(0, 0).text()) # cargo tecnico (team leader)
        sheet["C55"].value = "RHELEC" # Empresa
        sheet["C56"].value = self.tableTecnicosCuadrilla.item(0, 1).text() # telefono team leader
        sheet["C57"].value = self.tableTecnicosCuadrilla.item(0, 3).text() # correo team leader
        sheet["G53"].value = self.cbOyMs.currentText() # nombre O&M
        sheet["G54"].value = self.getTelefonoCorreoyCargoOyM()[2] # cargo O&M
        sheet["G56"].value = self.getTelefonoCorreoyCargoOyM()[0] # telefono O&M
        sheet["G57"].value = self.getTelefonoCorreoyCargoOyM()[1] # correo O&M
        pathOutputFile = "Output/"
        nameOutputFile = "SOLICITUD INGRESO PTI - " + self.cbRBS.currentText() + ".xlsx"
        book.save(pathOutputFile+nameOutputFile)
        book.close
        return pathOutputFile + nameOutputFile
    def getRegionZonaRBS(self):
        region = ""
        zona = ""
        with open("data/RBS.csv") as self.csvRBS:
            self.InfoRBS = reader(self.csvRBS)
            for self.row in self.InfoRBS:
                if str(self.row[0]).split(";")[1] == self.cbRBS.currentText():
                    return [str(self.row[0]).split(";")[2], str(self.row[0]).split(";")[3]]
        return [region, zona]
    def getEmailsFromRegionZona(self):
        r1z1uio = "tecomreddeaccesor1z1@claro.com.ec; nohemi.guadalupe@rhelec.ec; kevin.moreno@rhelec.ec; Azambrano@claro.com.ec; ernesto.quimbita@rhelec.ec; daniel.guisha@rhelec.ec; rodriguezlem@globalhitss.com; mariajose.orozco@rhelec.ec;maritza.nunez@rhelec.ec"
        r1z1prov = "tecomreddeaccesor1z1@claro.com.ec; Azambrano@claro.com.ec; nohemi.guadalupe@rhelec.ec; kevin.moreno@rhelec.ec; ernesto.quimbita@rhelec.ec; daniel.guisha@rhelec.ec; rodriguezlem@globalhitss.com; mariajose.orozco@rhelec.ec; maritza.nunez@rhelec.ec"
        r1z2uio = "tecoymrar1z2@claro.com.ec; lgarcia@claro.com.ec; nohemi.guadalupe@rhelec.ec; kevin.moreno@rhelec.ec; Azambrano@claro.com.ec; ernesto.quimbita@rhelec.ec; daniel.guisha@rhelec.ec; rodriguezlem@globalhitss.com; mariajose.orozco@rhelec.ec; maritza.nunez@rhelec.ec"
        r1z2prov = "tecoymrar1z2@claro.com.ec;  lgarcia@claro.com.ec; nohemi.guadalupe@rhelec.ec; kevin.moreno@rhelec.ec; Azambrano@claro.com.ec; ernesto.quimbita@rhelec.ec; daniel.guisha@rhelec.ec; rodriguezlem@globalhitss.com; mariajose.orozco@rhelec.ec; maritza.nunez@rhelec.ec"
        r2z1gye = "tecoymrar2z1@claro.com.ec; Azambrano@claro.com.ec; nohemi.guadalupe@rhelec.ec; kevin.moreno@rhelec.ec; alexander.santander@rhelec.ec; stalin.tutiven@rhelec.ec; christian.navas@rhelec.ec; rodriguezlem@globalhitss.com; mariajose.orozco@rhelec.ec"
        r2z1prov = "tecoymrar2z1@claro.com.ec; Azambrano@claro.com.ec; nohemi.guadalupe@rhelec.ec; kevin.moreno@rhelec.ec; alexander.santander@rhelec.ec; stalin.tutiven@rhelec.ec; rodriguezlem@globalhitss.com;  christian.navas@rhelec.ec; mariajose.orozco@rhelec.ec"
        r2z2gye = "tecoymrar2z2@claro.com.ec; Azambrano@claro.com.ec; nohemi.guadalupe@rhelec.ec; kevin.moreno@rhelec.ec; alexander.santander@rhelec.ec; stalin.tutiven@rhelec.ec; christian.navas@rhelec.ec; rodriguezlem@globalhitss.com; mariajose.orozco@rhelec.ec"
        r2z2prov = "tecoymrar2z2@claro.com.ec; azambrano@claro.com.ec; nohemi.guadalupe@rhelec.ec; kevin.moreno@rhelec.ec; alexander.santander@rhelec.ec;  stalin.tutiven@rhelec.ec; christian.navas@rhelec.ec; rodriguezlem@globalhitss.com; mariajose.orozco@rhelec.ec"
        if self.getRegionZonaRBS()[0] == "1" and self.getRegionZonaRBS()[1] == "1" and self.cbCuadrillas.currentText()[0:3] == "UIO":
            return r1z1uio
        elif self.getRegionZonaRBS()[0] == "1" and self.getRegionZonaRBS()[1] == "1" and self.cbCuadrillas.currentText()[0:3] != "UIO":
            return r1z1prov
        elif self.getRegionZonaRBS()[0] == "1" and self.getRegionZonaRBS()[1] == "2" and self.cbCuadrillas.currentText()[0:3] == "UIO":
            return r1z2uio
        elif self.getRegionZonaRBS()[0] == "1" and self.getRegionZonaRBS()[1] == "2" and self.cbCuadrillas.currentText()[0:3] != "UIO":
            return r1z2prov
        elif self.getRegionZonaRBS()[0] == "2" and self.getRegionZonaRBS()[1] == "1" and self.cbCuadrillas.currentText()[0:3] == "GYE":
            return r2z1gye
        elif self.getRegionZonaRBS()[0] == "2" and self.getRegionZonaRBS()[1] == "1" and self.cbCuadrillas.currentText()[0:3] != "GYE":
            return r2z1prov
        elif self.getRegionZonaRBS()[0] == "2" and self.getRegionZonaRBS()[1] == "2" and self.cbCuadrillas.currentText()[0:3] == "GYE":
            return r2z2gye
        elif self.getRegionZonaRBS()[0] == "2" and self.getRegionZonaRBS()[1] == "2" and self.cbCuadrillas.currentText()[0:3] != "GYE":
            return r2z2prov
        else: pass
    def makePermisos(self):
        self.pathPermisos = []
        for i in self.permisos.keys():
            if str(i).upper() == "FINANCIEROS" or str(i).upper() == "PRONTO":
                self.pathPermisos.append(self.makePermisoFinanciero(str(i)))
            elif str(i).upper() == "SBA":
                self.pathPermisos.append(self.makePermisoSBA())
            elif str(i).upper() == "ETC":
                self.pathPermisos.append(self.makePermisoETC())
            elif str(i).upper() == "PTI":
                self.pathPermisos.append(self.makePermisoPTI())
            elif str(i).upper() == "ATB":
                self.pathPermisos.append(self.makePermisoATB())
            elif str(i).upper() == "ECUSITES":
                self.pathPermisos.append(self.makePermisoECUSITES())
            elif str(i).upper() == "CNT":
                self.pathPermisos.append(self.makePermisoCNT())
            elif str(i).upper() == "TELEFONICA":
                self.pathPermisos.append(self.makePermisoTELEFONICA())
            else: pass
    def deleteOutputFolder(self):
        rmtree(getcwd()+"/Output")
        mkdir(getcwd()+"/Output")
    def openOutputFolder(self):
        startfile(getcwd()+"/Output")
    def cleanScreen(self):
        self.cbRBS.setCurrentIndex(0)
        self.lineAlarma.setText("")
        self.fixHeaderTables()
    def registerPermiso(self):
        global usuario
        fecha = datetime.today()
        for i in range(self.tablePermisosNecesarios.rowCount()):
            new_row = {"Usuario":[usuario], "Permiso":[self.tablePermisosNecesarios.item(i, 0).text()], 
                    "RBS":[self.cbRBS.currentText()], "Alarma":[self.lineAlarma.text()],
                    "Region":[self.lineRegion.text()], 
                    "Zona":[self.lineZona.text()], "Provincia":[self.lineProvincia.text()], 
                    "Canton":[self.getProvinciaCantonRBS()[1]], "O&M Solicitado":[self.cbOyMs.currentText()], 
                    "Supervisor":[self.cbSupervisores.currentText()], "Cuadrilla":[self.cbCuadrillas.currentText()], 
                    "Fecha Solicitada":[str(fecha.day)+"/"+str(fecha.month)+"/"+str(fecha.year)], 
                    "Hora Solicitada":[str(fecha.time())]}
            df_new = DataFrame(new_row)
            try:
                df_existing = read_excel(getcwd()+"/data/Registro Permisos.xlsx")
                df_combined = df_existing._append(df_new, ignore_index = True)
                df_combined.to_excel(getcwd()+"/data/Registro Permisos.xlsx", index = False)
            except:
                df_new.to_excel(getcwd()+"/data/Registro Permisos.xlsx", index = False)
    def enviarFormato(self):
        self.verificationInformation()
        if self.cbEliminarFormatos.isChecked():
            try:self.deleteOutputFolder()
            except:pass
        else: pass
        self.permisos = {}
        self.tecnicos = []
        if self.error == []:
            for i in range(self.tablePermisosNecesarios.rowCount()):
                self.permisos[self.tablePermisosNecesarios.item(i, 0).text()]  = [self.tablePermisosNecesarios.item(i, 1).text(), 
                                                                            self.tablePermisosNecesarios.item(i, 2).text(),
                                                                            self.tablePermisosNecesarios.item(i, 3).text(),
                                                                            self.tablePermisosNecesarios.item(i, 4).text()]
            self.makePermisos()
            print(self.pathPermisos)
            for i in range(self.tableTecnicosCuadrilla.rowCount()):
                if self.tableTecnicosCuadrilla.item(i, 0) != "":
                    self.tecnicos.append([self.tableTecnicosCuadrilla.item(i, 0), self.tableTecnicosCuadrilla.item(i, 1),
                                          self.tableTecnicosCuadrilla.item(i, 2), self.tableTecnicosCuadrilla.item(i, 3)])
            allData = {}
            allData["NameRBS"] = self.cbRBS.currentText()
            allData["Region"] = self.getRegionZonaRBS()[0]
            allData["Zona"] = self.getRegionZonaRBS()[1]
            allData["AllEmails"] = self.getEmailsFromRegionZona()
            allData["OyM"] = [self.cbOyMs.currentText(), self.getTelefonoCorreoyCargoOyM()[0], self.getTelefonoCorreoyCargoOyM()[1]]
            allData["Supervisor"] = [self.cbSupervisores.currentText(), self.getTelefonoyCorreoSupervisor()[0], self.getTelefonoyCorreoSupervisor()[1]]
            allData["Cuadrilla"] = self.cbCuadrillas.currentText()
            allData["Permisos"] = self.permisos
            allData["Tecnicos"] = self.tecnicos
            allData["Alarma"] = self.lineAlarma.text()
            self.registerPermiso()
            self.openOutputFolder()
            self.cleanScreen()

            #self.w = WindowSendEmail(allData = allData)
            #self.w.show()
            #self.close()
        else:
            print(self.error)

class Video_Logo_Window(QtWidgets.QMainWindow, QT_Windows.VideoLogoWindow_ui.Ui_Video_Logo_Window):
    def __init__(self):
        super(Video_Logo_Window, self).__init__()
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.FramelessWindowHint)
        #self.setAttribute(qtc.Qt.WA_TranslucentBackground)
        self.player = QtMultimedia.QMediaPlayer()
        self.player.setSource("ResourcesFolder/Logo_Rhelec/Logo_Rhelec.mp4")
        self.player.setVideoOutput(self.VideoPlace)
        self.VideoPlace.show()
        self.player.play()
        self.player.playbackStateChanged.connect(self.paroVideo)
    def convertXLSX2CSV(self):
        try:
            decrypted = BytesIO()
            with open("data/Informacion_RBS_Cuadrillas.xlsx", "rb") as f:
                data = msoffcrypto.OfficeFile(f)
                # Default passwords for encrypted excel sheets
                # Add your password here if it differs than the default
                data.load_key(password="RHELEC.2024")
                data.decrypt(decrypted)
            #Login
            df = read_excel(decrypted, sheet_name='Login')
            df.to_csv("data/Login.csv", index = None, header = True, sep=";", encoding='latin1')
            # Cuadrillas
            df = read_excel(decrypted, sheet_name='Cuadrillas')
            df.to_csv("data/Cuadrillas.csv", index = None, header = True, sep=";", encoding='latin1')
            # OyMs
            df = read_excel(decrypted, sheet_name='OyMs')
            df.to_csv("data/OyMs.csv", index = None, header = True, sep=";", encoding='latin1')
            # RBS
            df = read_excel(decrypted, sheet_name='RBS')
            df.to_csv("data/RBS.csv", index = None, header = True, sep=";", encoding='latin1')
            # Supervisores
            df = read_excel(decrypted, sheet_name='Supervisores')
            df.to_csv("data/Supervisores.csv", index = None, header = True, sep=";", encoding='latin1')
            return True
        except Exception as e:
            print(e)
            return False
    def paroVideo(self):
        if self.player.isPlaying():
            pass
        else:
            self.close()
            if self.convertXLSX2CSV():
                self.w = WindowLogin()
                self.w.show()
            else: print("No funca")

def deleteAllData():
    try:
        remove(getcwd()+"/Data/Cuadrillas.csv")
        remove(getcwd()+"/Data/Login.csv")
        remove(getcwd()+"/Data/OyMs.csv")
        remove(getcwd()+"/Data/RBS.csv")
        remove(getcwd()+"/Data/Supervisores.csv")
    except: pass

if __name__ == "__main__":
    if isfile(getcwd()+"/ResourcesFolder/Logo_Rhelec/Logo_Rhelec.mp4"):
        app = QtWidgets.QApplication(argv)
        window = Video_Logo_Window()
        window.show()
        app.exec()
    else:
        pass
        #window = WindowLogin()